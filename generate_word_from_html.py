import os
from bs4 import BeautifulSoup
from docx import Document

# Read HTML file
with open('dma.html', 'r', encoding='utf-8') as f:
    html = f.read()

soup = BeautifulSoup(html, 'html.parser')

doc = Document()
doc.add_heading('DisasterMap AI - Project Synopsis', 0)

# Map HTML heading tags to Word heading levels
tag_to_level = {
    'h1': 1,
    'h2': 2,
    'h3': 3,
    'h4': 4
}

# Find all headings in order
for tag in soup.find_all(['h1', 'h2', 'h3', 'h4']):
    text = tag.get_text(strip=True)
    level = tag_to_level[tag.name]
    doc.add_heading(text, level=level)

# Save the document
output_path = 'dma_headings.docx'
doc.save(output_path)
print(f'Word file created: {output_path}') 